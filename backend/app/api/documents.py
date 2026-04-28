from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Query
from fastapi.responses import StreamingResponse
from sqlmodel import Session, select
from pydantic import BaseModel
from typing import Optional
from datetime import date, datetime
import json
import io
import os

from ..database import get_session, ParsedDocument
from .auth import get_current_user, User
from ..services.doc_parser import parse_document_with_ai

router = APIRouter()

PACK_NUMBER = os.getenv("PACK_NUMBER", "44")
PACK_LOCATION = os.getenv("PACK_LOCATION", "Philipsburg, PA")

UPLOADS_DIR = "/data/documents/uploads"
HANDOUTS_DIR = "/data/documents/handouts"


class DocumentPatch(BaseModel):
    event_name: Optional[str] = None
    event_type: Optional[str] = None
    start_date: Optional[date] = None
    end_date: Optional[date] = None
    registration_deadline: Optional[date] = None
    location: Optional[str] = None
    address: Optional[str] = None
    cost_scout: Optional[str] = None
    cost_adult: Optional[str] = None
    cost_notes: Optional[str] = None
    contact_name: Optional[str] = None
    contact_email: Optional[str] = None
    contact_phone: Optional[str] = None
    registration_url: Optional[str] = None
    age_requirements: Optional[str] = None
    what_to_bring: Optional[str] = None
    key_notes: Optional[str] = None
    family_summary: Optional[str] = None


def _extract_text_from_pdf(data: bytes) -> str:
    import pdfplumber
    text_parts = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for page in pdf.pages:
            t = page.extract_text()
            if t:
                text_parts.append(t)
    return "\n".join(text_parts)


def _extract_text_from_image(data: bytes) -> str:
    import pytesseract
    from PIL import Image
    img = Image.open(io.BytesIO(data))
    return pytesseract.image_to_string(img)


def _parse_date_field(value) -> Optional[date]:
    if not value:
        return None
    if isinstance(value, date):
        return value
    try:
        return date.fromisoformat(str(value))
    except Exception:
        return None


def _safe_filename(doc_id: int, original: str) -> str:
    safe = "".join(c for c in original if c.isalnum() or c in "._-")[:80]
    return f"{doc_id}_{safe}"


@router.post("/parse", status_code=201)
async def parse_document(
    file: UploadFile = File(...),
    session: Session = Depends(get_session),
    current_user: User = Depends(get_current_user),
):
    filename = file.filename or "upload"
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    if ext not in {"pdf", "jpg", "jpeg", "png"}:
        raise HTTPException(400, "Only PDF, JPG, and PNG files are supported")

    data = await file.read()

    if ext == "pdf":
        raw_text = _extract_text_from_pdf(data)
    else:
        raw_text = _extract_text_from_image(data)

    if len(raw_text.strip()) < 100:
        raise HTTPException(400, "Could not extract enough text from the document")

    try:
        parsed = parse_document_with_ai(raw_text)
    except ValueError as e:
        raise HTTPException(503, str(e))
    except Exception as e:
        raise HTTPException(500, f"AI parsing failed: {e}")

    doc = ParsedDocument(
        filename=filename,
        raw_text=raw_text[:20000],
        uploaded_by=current_user.username,
        event_name=parsed.get("event_name"),
        event_type=parsed.get("event_type"),
        start_date=_parse_date_field(parsed.get("start_date")),
        end_date=_parse_date_field(parsed.get("end_date")),
        registration_deadline=_parse_date_field(parsed.get("registration_deadline")),
        location=parsed.get("location"),
        address=parsed.get("address"),
        cost_scout=parsed.get("cost_scout"),
        cost_adult=parsed.get("cost_adult"),
        cost_notes=parsed.get("cost_notes"),
        contact_name=parsed.get("contact_name"),
        contact_email=parsed.get("contact_email"),
        contact_phone=parsed.get("contact_phone"),
        registration_url=parsed.get("registration_url"),
        age_requirements=parsed.get("age_requirements"),
        what_to_bring=parsed.get("what_to_bring"),
        key_notes=json.dumps(parsed.get("key_notes") or []),
        family_summary=parsed.get("family_summary"),
    )
    session.add(doc)
    session.commit()
    session.refresh(doc)

    # Save original file to storage
    try:
        os.makedirs(UPLOADS_DIR, exist_ok=True)
        save_path = os.path.join(UPLOADS_DIR, _safe_filename(doc.id, filename))
        with open(save_path, "wb") as f:
            f.write(data)
        doc.source_file_path = save_path
        session.add(doc)
        session.commit()
        session.refresh(doc)
    except Exception:
        pass  # file save failure is non-fatal

    return doc


@router.get("/")
def list_documents(session: Session = Depends(get_session), _: User = Depends(get_current_user)):
    docs = session.exec(
        select(ParsedDocument).order_by(ParsedDocument.uploaded_at.desc())
    ).all()
    return [
        {
            "id": d.id,
            "filename": d.filename,
            "event_name": d.event_name,
            "event_type": d.event_type,
            "start_date": d.start_date,
            "uploaded_at": d.uploaded_at,
            "source_file_path": d.source_file_path,
            "handout_file_path": d.handout_file_path,
        }
        for d in docs
    ]


@router.get("/{doc_id}")
def get_document(doc_id: int, session: Session = Depends(get_session), _: User = Depends(get_current_user)):
    doc = session.get(ParsedDocument, doc_id)
    if not doc:
        raise HTTPException(404, "Document not found")
    return doc


@router.patch("/{doc_id}")
def patch_document(doc_id: int, patch: DocumentPatch, session: Session = Depends(get_session), _: User = Depends(get_current_user)):
    doc = session.get(ParsedDocument, doc_id)
    if not doc:
        raise HTTPException(404, "Document not found")
    for field, value in patch.model_dump(exclude_unset=True).items():
        setattr(doc, field, value)
    session.add(doc)
    session.commit()
    session.refresh(doc)
    return doc


@router.get("/{doc_id}/handout")
def download_handout(
    doc_id: int,
    regenerate: bool = Query(default=False),
    session: Session = Depends(get_session),
    _: User = Depends(get_current_user),
):
    doc = session.get(ParsedDocument, doc_id)
    if not doc:
        raise HTTPException(404, "Document not found")

    # Serve cached handout if available and not regenerating
    if not regenerate and doc.handout_file_path and os.path.isfile(doc.handout_file_path):
        with open(doc.handout_file_path, "rb") as f:
            buf = io.BytesIO(f.read())
        buf.seek(0)
        safe_name = (doc.event_name or "handout").replace(" ", "_")[:40]
        date_str = str(doc.start_date) if doc.start_date else "undated"
        filename = f"{safe_name}_{date_str}_Handout.docx"
        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )

    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise HTTPException(500, "python-docx not installed")

    d = Document()

    # ── Header ────────────────────────────────────────────────────────────────
    title_para = d.add_heading(doc.event_name or "Event Information", 0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title_para.runs:
        run.font.color.rgb = RGBColor(0x00, 0x30, 0x87)  # navy

    if doc.event_type:
        sub = d.add_paragraph(doc.event_type)
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub.runs[0].italic = True

    d.add_paragraph("")

    # ── Info table ─────────────────────────────────────────────────────────────
    def fmt_dates(s, e):
        if s and e:
            return f"{s.strftime('%B %d')} – {e.strftime('%B %d, %Y')}"
        return s.strftime("%B %d, %Y") if s else ""

    rows = [
        ("Dates",                  fmt_dates(doc.start_date, doc.end_date)),
        ("Location",               doc.location),
        ("Address",                doc.address),
        ("Scout Cost",             doc.cost_scout),
        ("Adult Cost",             doc.cost_adult),
        ("Cost Notes",             doc.cost_notes),
        ("Registration Deadline",  doc.registration_deadline.strftime("%B %d, %Y") if doc.registration_deadline else None),
        ("Age Requirements",       doc.age_requirements),
        ("Contact",                f"{doc.contact_name or ''} {doc.contact_email or ''} {doc.contact_phone or ''}".strip() or None),
        ("Register At",            doc.registration_url),
    ]
    rows = [(label, val) for label, val in rows if val]

    if rows:
        table = d.add_table(rows=len(rows), cols=2)
        table.style = "Table Grid"
        for i, (label, val) in enumerate(rows):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = val or ""
            table.rows[i].cells[0].paragraphs[0].runs[0].bold = True

    d.add_paragraph("")

    # ── Key notes ──────────────────────────────────────────────────────────────
    notes = []
    if doc.key_notes:
        try:
            notes = json.loads(doc.key_notes)
        except Exception:
            pass
    if notes:
        d.add_heading("What to Know", 2)
        for note in notes:
            if note:
                d.add_paragraph(str(note), style="List Bullet")

    # ── What to bring ──────────────────────────────────────────────────────────
    if doc.what_to_bring:
        d.add_heading("What to Bring", 2)
        d.add_paragraph(doc.what_to_bring)

    # ── Family summary ─────────────────────────────────────────────────────────
    if doc.family_summary:
        d.add_paragraph("")
        summary_para = d.add_paragraph(doc.family_summary)
        for run in summary_para.runs:
            run.italic = True

    # ── Footer ─────────────────────────────────────────────────────────────────
    d.add_paragraph("")
    footer = d.add_paragraph(f"Pack {PACK_NUMBER} · {PACK_LOCATION}")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    footer.runs[0].font.size = Pt(9)

    buf = io.BytesIO()
    d.save(buf)
    buf.seek(0)

    safe_name = (doc.event_name or "handout").replace(" ", "_")[:40]
    date_str = str(doc.start_date) if doc.start_date else "undated"
    filename = f"{safe_name}_{date_str}_Handout.docx"

    # Save handout to disk
    try:
        os.makedirs(HANDOUTS_DIR, exist_ok=True)
        handout_path = os.path.join(HANDOUTS_DIR, f"{doc.id}_{safe_name}_Handout.docx")
        with open(handout_path, "wb") as f:
            f.write(buf.getvalue())
        doc.handout_file_path = handout_path
        session.add(doc)
        session.commit()
    except Exception:
        pass

    buf.seek(0)
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.get("/{doc_id}/source")
def download_source(doc_id: int, session: Session = Depends(get_session), _: User = Depends(get_current_user)):
    doc = session.get(ParsedDocument, doc_id)
    if not doc:
        raise HTTPException(404, "Document not found")
    if not doc.source_file_path or not os.path.isfile(doc.source_file_path):
        raise HTTPException(404, "Source file not available")

    ext = doc.filename.rsplit(".", 1)[-1].lower() if "." in doc.filename else "bin"
    media_types = {
        "pdf": "application/pdf",
        "jpg": "image/jpeg",
        "jpeg": "image/jpeg",
        "png": "image/png",
    }
    media_type = media_types.get(ext, "application/octet-stream")

    with open(doc.source_file_path, "rb") as f:
        data = f.read()

    return StreamingResponse(
        io.BytesIO(data),
        media_type=media_type,
        headers={"Content-Disposition": f'attachment; filename="{doc.filename}"'},
    )


@router.delete("/{doc_id}", status_code=204)
def delete_document(doc_id: int, session: Session = Depends(get_session), _: User = Depends(get_current_user)):
    doc = session.get(ParsedDocument, doc_id)
    if not doc:
        raise HTTPException(404, "Document not found")
    # Delete associated files
    for path in [doc.source_file_path, doc.handout_file_path]:
        if path and os.path.isfile(path):
            try:
                os.remove(path)
            except Exception:
                pass
    session.delete(doc)
    session.commit()
