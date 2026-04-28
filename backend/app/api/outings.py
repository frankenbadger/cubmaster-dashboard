from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from sqlmodel import Session, select
from pydantic import BaseModel
from typing import Optional
from datetime import date, datetime
import json
import io
import os

from ..database import get_session, Outing
from .auth import get_current_user, User

router = APIRouter()

PACK_NUMBER = os.getenv("PACK_NUMBER", "44")
PACK_LOCATION = os.getenv("PACK_LOCATION", "Philipsburg, PA")

DEFAULT_CHECKLIST = [
    {"label": "Choose and confirm location", "done": False},
    {"label": "Get council/charter org approval if needed", "done": False},
    {"label": "Set date and time", "done": False},
    {"label": "Create permission slip", "done": False},
    {"label": "Send info to families (Band post)", "done": False},
    {"label": "Collect permission slips", "done": False},
    {"label": "Confirm headcount", "done": False},
    {"label": "Arrange transportation", "done": False},
    {"label": "First aid kit assigned", "done": False},
    {"label": "Emergency contact list printed", "done": False},
]


class OutingCreate(BaseModel):
    name: str
    outing_type: str
    status: str = "planning"
    date_start: Optional[date] = None
    date_end: Optional[date] = None
    location_name: Optional[str] = None
    location_address: Optional[str] = None
    meeting_time: Optional[str] = None
    return_time: Optional[str] = None
    cost_scout: Optional[str] = None
    cost_adult: Optional[str] = None
    cost_notes: Optional[str] = None
    max_participants: Optional[int] = None
    min_participants: Optional[int] = None
    transportation: Optional[str] = None
    gear_needed: Optional[str] = None
    permission_slip_needed: bool = True
    permission_slip_notes: Optional[str] = None
    medical_form_needed: bool = False
    contact_name: Optional[str] = None
    contact_phone: Optional[str] = None
    contact_email: Optional[str] = None
    reservation_url: Optional[str] = None
    reservation_confirmation: Optional[str] = None
    notes: Optional[str] = None
    checklist: Optional[str] = None


class OutingPatch(BaseModel):
    name: Optional[str] = None
    outing_type: Optional[str] = None
    status: Optional[str] = None
    date_start: Optional[date] = None
    date_end: Optional[date] = None
    location_name: Optional[str] = None
    location_address: Optional[str] = None
    meeting_time: Optional[str] = None
    return_time: Optional[str] = None
    cost_scout: Optional[str] = None
    cost_adult: Optional[str] = None
    cost_notes: Optional[str] = None
    max_participants: Optional[int] = None
    min_participants: Optional[int] = None
    transportation: Optional[str] = None
    gear_needed: Optional[str] = None
    permission_slip_needed: Optional[bool] = None
    permission_slip_notes: Optional[str] = None
    medical_form_needed: Optional[bool] = None
    contact_name: Optional[str] = None
    contact_phone: Optional[str] = None
    contact_email: Optional[str] = None
    reservation_url: Optional[str] = None
    reservation_confirmation: Optional[str] = None
    notes: Optional[str] = None
    checklist: Optional[str] = None


@router.get("/")
def list_outings(session: Session = Depends(get_session), _: User = Depends(get_current_user)):
    return session.exec(
        select(Outing).order_by(Outing.date_start.asc().nullslast(), Outing.created_at.desc())
    ).all()


@router.get("/{outing_id}")
def get_outing(outing_id: int, session: Session = Depends(get_session), _: User = Depends(get_current_user)):
    outing = session.get(Outing, outing_id)
    if not outing:
        raise HTTPException(404, "Outing not found")
    return outing


@router.post("/", status_code=201)
def create_outing(data: OutingCreate, session: Session = Depends(get_session), current_user: User = Depends(get_current_user)):
    checklist = data.checklist or json.dumps(DEFAULT_CHECKLIST)
    outing = Outing(
        **{k: v for k, v in data.model_dump().items() if k != "checklist"},
        checklist=checklist,
        created_by=current_user.username,
    )
    session.add(outing)
    session.commit()
    session.refresh(outing)
    return outing


@router.patch("/{outing_id}")
def patch_outing(outing_id: int, patch: OutingPatch, session: Session = Depends(get_session), _: User = Depends(get_current_user)):
    outing = session.get(Outing, outing_id)
    if not outing:
        raise HTTPException(404, "Outing not found")
    for field, value in patch.model_dump(exclude_unset=True).items():
        setattr(outing, field, value)
    outing.updated_at = datetime.utcnow()
    session.add(outing)
    session.commit()
    session.refresh(outing)
    return outing


@router.delete("/{outing_id}", status_code=204)
def delete_outing(outing_id: int, session: Session = Depends(get_session), _: User = Depends(get_current_user)):
    outing = session.get(Outing, outing_id)
    if not outing:
        raise HTTPException(404, "Outing not found")
    session.delete(outing)
    session.commit()


@router.get("/{outing_id}/signup-genius-url")
def signup_genius_url(outing_id: int, session: Session = Depends(get_session), _: User = Depends(get_current_user)):
    outing = session.get(Outing, outing_id)
    if not outing:
        raise HTTPException(404, "Outing not found")

    date_str = str(outing.date_start) if outing.date_start else "TBD"
    end_str = str(outing.date_end) if outing.date_end else date_str
    location = outing.location_name or ""
    if outing.location_address:
        location += f", {outing.location_address}"

    lines = [f"Title: Pack {PACK_NUMBER} – {outing.name}"]
    if outing.date_start:
        lines.append(f"Date: {outing.date_start.strftime('%B %d, %Y')}")
    if outing.date_end and outing.date_end != outing.date_start:
        lines.append(f"End Date: {outing.date_end.strftime('%B %d, %Y')}")
    if outing.meeting_time:
        lines.append(f"Start Time: {outing.meeting_time}")
    if outing.return_time:
        lines.append(f"End Time: {outing.return_time}")
    if location:
        lines.append(f"Location: {location}")
    if outing.cost_scout:
        lines.append(f"Cost per Scout: {outing.cost_scout}")
    if outing.cost_adult:
        lines.append(f"Cost per Adult: {outing.cost_adult}")
    if outing.notes:
        lines.append(f"Notes: {outing.notes}")

    return {
        "signup_genius_url": "https://www.signupgenius.com/CreateSignUp.cfm",
        "prefill_text": "\n".join(lines),
    }


@router.get("/{outing_id}/handout")
def download_handout(outing_id: int, session: Session = Depends(get_session), _: User = Depends(get_current_user)):
    outing = session.get(Outing, outing_id)
    if not outing:
        raise HTTPException(404, "Outing not found")

    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise HTTPException(500, "python-docx not installed")

    d = Document()

    # ── Header ────────────────────────────────────────────────────────────────
    title_para = d.add_heading(outing.name, 0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title_para.runs:
        run.font.color.rgb = RGBColor(0x00, 0x30, 0x87)

    type_sub = d.add_paragraph(f"{outing.outing_type} · Pack {PACK_NUMBER}")
    type_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    type_sub.runs[0].italic = True

    d.add_paragraph("")

    # ── Info table ─────────────────────────────────────────────────────────────
    def fmt_date(dt):
        return dt.strftime("%B %d, %Y") if dt else ""

    def fmt_dates(s, e):
        if s and e and s != e:
            return f"{s.strftime('%B %d')} – {e.strftime('%B %d, %Y')}"
        return fmt_date(s)

    location = outing.location_name or ""
    if outing.location_address:
        location += f"\n{outing.location_address}" if location else outing.location_address

    rows = [
        ("Date",              fmt_dates(outing.date_start, outing.date_end)),
        ("Location",          outing.location_name),
        ("Address",           outing.location_address),
        ("Meeting Time",      outing.meeting_time),
        ("Return Time",       outing.return_time),
        ("Scout Cost",        outing.cost_scout),
        ("Adult Cost",        outing.cost_adult),
        ("Cost Notes",        outing.cost_notes),
        ("Max Participants",  str(outing.max_participants) if outing.max_participants else None),
        ("Transportation",    outing.transportation),
        ("Contact",           outing.contact_name),
        ("Phone",             outing.contact_phone),
        ("Email",             outing.contact_email),
        ("Reservation URL",   outing.reservation_url),
        ("Confirmation #",    outing.reservation_confirmation),
    ]
    rows = [(label, val) for label, val in rows if val]

    if rows:
        table = d.add_table(rows=len(rows), cols=2)
        table.style = "Table Grid"
        for i, (label, val) in enumerate(rows):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = val
            table.rows[i].cells[0].paragraphs[0].runs[0].bold = True

    d.add_paragraph("")

    # ── Gear list ──────────────────────────────────────────────────────────────
    gear = []
    if outing.gear_needed:
        try:
            gear = json.loads(outing.gear_needed)
        except Exception:
            pass
    if gear:
        d.add_heading("What to Bring", 2)
        for item in gear:
            if item:
                d.add_paragraph(str(item), style="List Bullet")

    # ── Notes ──────────────────────────────────────────────────────────────────
    if outing.notes:
        d.add_heading("Additional Notes", 2)
        d.add_paragraph(outing.notes)

    # ── Permission slip tear-off ───────────────────────────────────────────────
    d.add_page_break()
    slip_title = d.add_heading("Permission Slip / Sign-Up", 1)
    slip_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    info = d.add_paragraph(
        f"Pack {PACK_NUMBER} – {outing.name}\n"
        f"Date: {fmt_dates(outing.date_start, outing.date_end)}\n"
        f"Location: {outing.location_name or 'TBD'}"
    )
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER

    d.add_paragraph("")
    d.add_paragraph("Please complete and return by: ___________________________")
    d.add_paragraph("")

    slip_rows = [
        ("Scout's Full Name",        ""),
        ("Den",                      ""),
        ("# of Scouts Attending",    ""),
        ("# of Adults Attending",    ""),
        ("Emergency Contact Name",   ""),
        ("Emergency Contact Phone",  ""),
        ("Medical Concerns / Allergies", ""),
    ]
    if outing.cost_scout:
        slip_rows.append((f"Payment Enclosed ({outing.cost_scout}/Scout)", "☐ Yes  ☐ Cash  ☐ Check"))
    if outing.medical_form_needed:
        slip_rows.append(("BSA Medical Form Attached", "☐ Yes"))

    slip_table = d.add_table(rows=len(slip_rows), cols=2)
    slip_table.style = "Table Grid"
    for i, (label, val) in enumerate(slip_rows):
        slip_table.rows[i].cells[0].text = label
        slip_table.rows[i].cells[1].text = val
        slip_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True

    d.add_paragraph("")
    sig = d.add_paragraph("Parent/Guardian Signature: _________________________________  Date: _____________")

    d.add_paragraph("")
    footer = d.add_paragraph(f"Pack {PACK_NUMBER} · {PACK_LOCATION}")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    footer.runs[0].font.size = Pt(9)

    buf = io.BytesIO()
    d.save(buf)
    buf.seek(0)

    safe_name = outing.name.replace(" ", "_")[:40]
    date_str = str(outing.date_start) if outing.date_start else "undated"
    filename = f"{safe_name}_Info_{date_str}.docx"

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
