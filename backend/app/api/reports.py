from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from sqlmodel import Session, select
from datetime import datetime
from pydantic import BaseModel
from typing import Optional
import json
import io
import os

from ..database import get_session, MonthlyReport
from .auth import get_current_user, User

router = APIRouter()

PACK_NUMBER = os.getenv("PACK_NUMBER", "44")
PACK_LOCATION = os.getenv("PACK_LOCATION", "Philipsburg, PA")

MONTH_NAMES = ["January","February","March","April","May","June",
               "July","August","September","October","November","December"]

DEN_ORDER = ['Lions', 'Tigers', 'Wolves', 'Bears', 'Webelos', 'AOL']

OPTIONAL_SECTION_HEADINGS = {
    '1.2': '1.2 Additional Meeting Notes',
    '1.3': '1.3 Training / Leader Development',
    '2.1': '2.1 Den Assignments',
    '3.1': '3.1 Council Events of Interest',
    '5.1': '5.1 Membership & Recruitment',
    '5.2': '5.2 Financial Notes',
}

OPTIONAL_SECTION_GROUP = {
    '1.2': 1, '1.3': 1,
    '2.1': 2,
    '3.1': 3,
    '5.1': 5, '5.2': 5,
}


class ReportUpsert(BaseModel):
    year: int
    month: int
    last_meeting_name: Optional[str] = None
    last_meeting_summary: Optional[str] = None
    last_meeting_attendance: Optional[str] = None
    last_meeting_went_well: Optional[str] = None
    last_meeting_needs_improvement: Optional[str] = None
    upcoming_meeting_program: Optional[str] = None
    upcoming_meeting_agenda: Optional[str] = None
    upcoming_events: Optional[str] = None
    den_updates: Optional[str] = None
    notes: Optional[str] = None
    potential_outings: Optional[str] = None


def _parse_notes_field(raw: Optional[str]):
    """Split __sections__JSON\\nplaintext into (added_keys, content_dict, plain_notes)."""
    if not raw or not raw.startswith('__sections__'):
        return [], {}, raw or ''
    rest = raw[len('__sections__'):]
    nl = rest.find('\n')
    json_str = rest[:nl] if nl != -1 else rest
    plain = rest[nl + 1:] if nl != -1 else ''
    try:
        data = json.loads(json_str)
        added = data.get('added', [])
        content = {k: v for k, v in data.items() if k != 'added'}
        return added, content, plain
    except Exception:
        return [], {}, raw


@router.get("/")
def list_reports(session: Session = Depends(get_session), _: User = Depends(get_current_user)):
    return session.exec(
        select(MonthlyReport).order_by(MonthlyReport.year.desc(), MonthlyReport.month.desc())
    ).all()


@router.get("/{year}/{month}")
def get_report(year: int, month: int, session: Session = Depends(get_session), _: User = Depends(get_current_user)):
    report = session.exec(
        select(MonthlyReport).where(MonthlyReport.year == year, MonthlyReport.month == month)
    ).first()
    return report or MonthlyReport(year=year, month=month)


@router.put("/{year}/{month}")
def upsert_report(year: int, month: int, data: ReportUpsert,
                  session: Session = Depends(get_session), current_user: User = Depends(get_current_user)):
    report = session.exec(
        select(MonthlyReport).where(MonthlyReport.year == year, MonthlyReport.month == month)
    ).first()
    if not report:
        report = MonthlyReport(year=year, month=month)
    for field, value in data.model_dump(exclude={"year", "month"}).items():
        if value is not None:
            setattr(report, field, value)
    report.updated_at = datetime.utcnow()
    report.updated_by = current_user.username
    session.add(report)
    session.commit()
    session.refresh(report)
    return report


@router.get("/{year}/{month}/download")
def download_report(year: int, month: int, session: Session = Depends(get_session),
                    _: User = Depends(get_current_user)):
    report = session.exec(
        select(MonthlyReport).where(MonthlyReport.year == year, MonthlyReport.month == month)
    ).first()

    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise HTTPException(status_code=500, detail="python-docx not installed")

    doc = Document()

    # Title block
    title = doc.add_heading("Cubmaster's Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph(f"Pack {PACK_NUMBER} · {PACK_LOCATION} · {MONTH_NAMES[month]} {year}")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Parse optional sections from the notes field
    added_keys, opt_content, plain_notes = _parse_notes_field(report.notes if report else None)

    def add_optional_sections(group: int):
        for key in added_keys:
            if OPTIONAL_SECTION_GROUP.get(key) == group:
                doc.add_heading(OPTIONAL_SECTION_HEADINGS.get(key, key), 2)
                text = opt_content.get(key, '')
                if text:
                    _add_bullets(doc, None, text)

    # ── Section 1 ─────────────────────────────────────────────────────────────
    doc.add_heading("1. Monthly Review", 1)

    h11 = "1.1 Last Pack Meeting"
    if report and report.last_meeting_name:
        h11 += f" – {report.last_meeting_name}"
    doc.add_heading(h11, 2)

    if report:
        _add_field(doc, "Summary", report.last_meeting_summary)
        _add_field(doc, "Attendance", report.last_meeting_attendance)
        _add_bullets(doc, "Went Well", report.last_meeting_went_well)
        _add_bullets(doc, "Needs Improvement", report.last_meeting_needs_improvement)

    add_optional_sections(1)

    # ── Section 2 ─────────────────────────────────────────────────────────────
    doc.add_heading("2. Upcoming Pack Meeting", 1)
    if report:
        _add_field(doc, "Program", report.upcoming_meeting_program)
        _add_field(doc, "Agenda", report.upcoming_meeting_agenda)

    add_optional_sections(2)

    # ── Section 3 ─────────────────────────────────────────────────────────────
    doc.add_heading("3. Upcoming Events", 1)
    if report and report.upcoming_events:
        _add_bullets(doc, None, report.upcoming_events)

    add_optional_sections(3)

    # ── Section 4 ─────────────────────────────────────────────────────────────
    doc.add_heading("4. Den Updates", 1)
    den_updates = {}
    if report and report.den_updates:
        try:
            den_updates = json.loads(report.den_updates)
        except Exception:
            pass
    for den_name in DEN_ORDER:
        text = den_updates.get(den_name, '').strip()
        if text:
            doc.add_heading(den_name, 3)
            _add_bullets(doc, None, text)

    # ── Section 5 ─────────────────────────────────────────────────────────────
    doc.add_heading("5. Notes", 1)
    if plain_notes:
        doc.add_paragraph(plain_notes)

    add_optional_sections(5)

    # ── Section 6 — Potential Outings ─────────────────────────────────────────
    outings = []
    if report and report.potential_outings:
        try:
            outings = json.loads(report.potential_outings)
        except Exception:
            pass

    if outings:
        doc.add_heading("6. Potential Outings", 1)
        for outing in outings:
            name = outing.get('name') or 'Unnamed Outing'
            doc.add_heading(name, 3)
            rows = [
                ("Type",            outing.get('type', '')),
                ("Timeframe",       outing.get('timeframe', '')),
                ("Est. Cost/Scout", outing.get('cost', '')),
                ("Notes",           outing.get('notes', '')),
            ]
            rows = [(label, val) for label, val in rows if val]
            if rows:
                table = doc.add_table(rows=len(rows), cols=2)
                table.style = "Table Grid"
                for i, (label, val) in enumerate(rows):
                    cell_label = table.rows[i].cells[0]
                    cell_val   = table.rows[i].cells[1]
                    cell_label.text = label
                    cell_val.text   = val
                    cell_label.paragraphs[0].runs[0].bold = True

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    filename = f"Cubmasters_Report_{MONTH_NAMES[month]}_{year}.docx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


def _add_field(doc, label, value):
    if not value:
        return
    p = doc.add_paragraph()
    run = p.add_run(f"{label}: ")
    run.bold = True
    p.add_run(value)


def _add_bullets(doc, label, value):
    if not value:
        return
    if label:
        p = doc.add_paragraph(label + ":")
        p.runs[0].bold = True
    for line in value.split("\n"):
        line = line.strip().lstrip("-•").strip()
        if line:
            doc.add_paragraph(line, style="List Bullet")
